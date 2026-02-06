#!/usr/bin/env python3
"""
Generate styled Word documents for job applications — Web Chat Edition.

Standalone script for Claude's Analysis sandbox. Duplicates core document
generation from generate_word_docs.py so it works without filesystem access
to the skill directory.

Also provides:
- generate_profile_yaml(): Export candidate data to portable YAML
- parse_uploaded_resume(): Extract structured text from uploaded .docx

Usage in Analysis tool:
    exec(open("generate_word_docs_web.py").read())

    generate_application_documents(
        candidate=candidate_data,
        job=job_data,
        cover_letter=cover_letter_content,
        resume=resume_content,
        output_dir="/tmp/"
    )
"""

import subprocess
import sys

# Auto-install python-docx if missing (Analysis sandbox has pip)
try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    import docx

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
from pathlib import Path
import re
import zipfile


# =============================================================================
# WCAG AA Accessible Colors (7:1+ contrast ratio on white)
# =============================================================================
ACCENT_COLOR = RGBColor(0x1a, 0x52, 0x76)  # Navy Blue #1a5276 - 8:1 contrast
TEXT_PRIMARY = RGBColor(0x1a, 0x20, 0x2c)  # Near-black #1a202c - 16:1 contrast
TEXT_SECONDARY = RGBColor(0x37, 0x41, 0x51)  # Dark Gray #374151 - 10:1 contrast

# Spacing constants
SPACE_BEFORE_HEADING = Pt(14)
SPACE_AFTER_HEADING = Pt(10)


# =============================================================================
# Helper Functions
# =============================================================================

def set_document_single_spacing(doc):
    """Set single line spacing as default for the entire document."""
    style = doc.styles['Normal']
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    if 'List Bullet' in doc.styles:
        list_style = doc.styles['List Bullet']
        list_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def add_horizontal_line(paragraph):
    """Add a colored horizontal line below a paragraph."""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:color'), '1a5276')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_section_heading(doc, text, is_first=False):
    """Add a styled section heading with proper spacing."""
    section = doc.add_paragraph()
    section.paragraph_format.space_before = Pt(6) if is_first else SPACE_BEFORE_HEADING
    section.paragraph_format.space_after = SPACE_AFTER_HEADING

    section_run = section.add_run(text)
    section_run.bold = True
    section_run.font.size = Pt(12)
    section_run.font.color.rgb = ACCENT_COLOR
    add_horizontal_line(section)

    return section


def sanitize_filename(text, max_length=50):
    """Convert text to a safe filename component."""
    safe = re.sub(r'[/\\:*?"<>|]', '', text)
    safe = re.sub(r'[-\s]+', '_', safe)
    safe = re.sub(r'[\x00-\x1f\x7f]', '', safe)
    safe = safe.strip('_')[:max_length].rstrip('_')
    return safe


def _add_text_with_highlights(paragraph, text, highlights):
    """Add text to paragraph with certain phrases bolded."""
    if not highlights:
        paragraph.add_run(text)
        return

    highlights = sorted(highlights, key=len, reverse=True)

    remaining = text
    for highlight in highlights:
        if highlight in remaining:
            parts = remaining.split(highlight, 1)
            if parts[0]:
                paragraph.add_run(parts[0])
            paragraph.add_run(highlight).bold = True
            remaining = parts[1] if len(parts) > 1 else ''

    if remaining:
        paragraph.add_run(remaining)


# =============================================================================
# Cover Letter Generator
# =============================================================================

def create_cover_letter(candidate, job, content, output_path):
    """
    Create a styled cover letter document.

    Args:
        candidate: dict with keys: name, phone, email, linkedin, calendar
        job: dict with keys: title, company, location, date (optional)
        content: dict with keys: opening, sections (list of {title, paragraphs}),
                 closing, signature_title
        output_path: Full path for output file
    """
    doc = Document()
    set_document_single_spacing(doc)

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # === HEADER ===
    name_p = doc.add_paragraph()
    name_run = name_p.add_run(candidate['name'])
    name_run.bold = True
    name_run.font.size = Pt(22)
    name_run.font.color.rgb = ACCENT_COLOR
    name_p.space_after = Pt(2)

    contact_parts = [candidate.get('phone', ''), candidate.get('email', ''),
                     candidate.get('linkedin', ''), candidate.get('calendar', '')]
    contact_text = ' | '.join(p for p in contact_parts if p)

    contact_p = doc.add_paragraph()
    contact_run = contact_p.add_run(contact_text)
    contact_run.font.size = Pt(10)
    contact_run.font.color.rgb = TEXT_SECONDARY
    add_horizontal_line(contact_p)
    contact_p.space_after = Pt(20)

    # === DATE & RECIPIENT ===
    date_str = job.get('date', datetime.now().strftime('%B %d, %Y'))
    doc.add_paragraph(date_str).space_after = Pt(16)

    recipient = doc.add_paragraph()
    recipient.add_run(f"Hiring Manager\n{job['company']}\n{job.get('location', '')}".strip())
    recipient.space_after = Pt(12)

    subject = doc.add_paragraph()
    subject_run = subject.add_run(f"Re: {job['title']}")
    subject_run.bold = True
    subject.space_after = Pt(16)

    doc.add_paragraph("Dear Hiring Manager,").space_after = Pt(10)

    # === OPENING PARAGRAPH ===
    opening = doc.add_paragraph()
    opening.add_run(content['opening'])
    opening.space_after = Pt(12)

    # === CONTENT SECTIONS ===
    for i, section in enumerate(content.get('sections', [])):
        add_section_heading(doc, section['title'], is_first=(i == 0))

        for para_content in section.get('paragraphs', []):
            if isinstance(para_content, dict):
                p = doc.add_paragraph()
                if para_content.get('label'):
                    p.add_run(f"{para_content['label']}: ").bold = True

                text = para_content.get('text', '')
                highlights = para_content.get('highlights', [])
                _add_text_with_highlights(p, text, highlights)
                p.space_after = Pt(8)

            elif isinstance(para_content, list):
                for bullet_item in para_content:
                    bullet = doc.add_paragraph(style='List Bullet')
                    if isinstance(bullet_item, dict):
                        bullet.add_run(f"{bullet_item['label']}: ").bold = True
                        bullet.add_run(bullet_item.get('text', ''))
                    else:
                        bullet.add_run(bullet_item)
                    bullet.space_after = Pt(4)

            else:
                p = doc.add_paragraph()
                p.add_run(str(para_content))
                p.space_after = Pt(8)

    # === CLOSING ===
    closing_text = content.get('closing', 'Best regards,')
    doc.add_paragraph(closing_text).space_after = Pt(20)

    sig_name = doc.add_paragraph()
    sig_name.add_run(candidate['name']).bold = True
    sig_name.space_after = Pt(2)

    if content.get('signature_title'):
        sig_title = doc.add_paragraph()
        sig_run = sig_title.add_run(content['signature_title'])
        sig_run.font.color.rgb = TEXT_SECONDARY
        sig_run.font.size = Pt(10)

    doc.save(output_path)
    return output_path


# =============================================================================
# Resume Generator
# =============================================================================

def create_resume(candidate, content, output_path):
    """
    Create a styled resume document.

    Args:
        candidate: dict with keys: name, title, phone, email, linkedin, calendar
        content: dict with keys: summary, skills (list of {category, items}),
                 experience (list of jobs), certifications, education
        output_path: Full path for output file
    """
    doc = Document()
    set_document_single_spacing(doc)

    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # === HEADER ===
    name_p = doc.add_paragraph()
    name_run = name_p.add_run(candidate['name'])
    name_run.bold = True
    name_run.font.size = Pt(24)
    name_run.font.color.rgb = ACCENT_COLOR
    name_p.space_after = Pt(2)

    if candidate.get('title'):
        title_p = doc.add_paragraph()
        title_run = title_p.add_run(candidate['title'])
        title_run.bold = True
        title_run.font.size = Pt(12)
        title_p.space_after = Pt(4)

    contact_parts = [candidate.get('phone', ''), candidate.get('email', ''),
                     candidate.get('linkedin', ''), candidate.get('calendar', '')]
    contact_text = ' | '.join(p for p in contact_parts if p)

    contact_p = doc.add_paragraph()
    contact_run = contact_p.add_run(contact_text)
    contact_run.font.size = Pt(10)
    contact_run.font.color.rgb = TEXT_SECONDARY
    contact_p.space_after = Pt(8)

    # === SUMMARY ===
    if content.get('summary'):
        add_section_heading(doc, 'Summary', is_first=True)
        summary_p = doc.add_paragraph()
        summary_p.add_run(content['summary'])
        summary_p.space_after = Pt(6)

    # === TECHNICAL SKILLS ===
    if content.get('skills'):
        add_section_heading(doc, 'Technical Skills')
        for skill in content['skills']:
            p = doc.add_paragraph()
            p.add_run(f"{skill['category']}:").bold = True
            p.add_run(f" {skill['items']}")
            p.space_after = Pt(4)
        doc.paragraphs[-1].space_after = Pt(6)

    # === PROFESSIONAL EXPERIENCE ===
    if content.get('experience'):
        add_section_heading(doc, 'Professional Experience')

        for idx, job in enumerate(content['experience']):
            job_title_p = doc.add_paragraph()
            job_title_p.paragraph_format.space_before = Pt(0) if idx == 0 else Pt(12)
            job_title_run = job_title_p.add_run(job['title'])
            job_title_run.bold = True
            job_title_run.font.size = Pt(11)
            job_title_p.space_after = Pt(1)

            company_p = doc.add_paragraph()
            company_run = company_p.add_run(job['company'])
            company_run.bold = True
            company_run.font.size = Pt(10)
            company_p.space_after = Pt(1)

            dates_p = doc.add_paragraph()
            dates_run = dates_p.add_run(job['dates'])
            dates_run.font.color.rgb = TEXT_SECONDARY
            dates_run.font.size = Pt(10)
            dates_p.space_after = Pt(4)

            for bullet in job.get('bullets', []):
                bullet_p = doc.add_paragraph(style='List Bullet')
                bullet_p.paragraph_format.left_indent = Inches(0.35)

                if isinstance(bullet, dict):
                    _add_text_with_highlights(bullet_p, bullet['text'],
                                              bullet.get('highlights', []))
                else:
                    bullet_p.add_run(bullet)

                bullet_p.space_after = Pt(2)

    # === CERTIFICATIONS ===
    if content.get('certifications'):
        add_section_heading(doc, 'Certifications')
        for cert in content['certifications']:
            cert_p = doc.add_paragraph(style='List Bullet')
            if isinstance(cert, dict):
                cert_p.add_run(cert['name']).bold = True
                cert_p.add_run(f" - {cert.get('detail', '')}")
            else:
                cert_p.add_run(cert)
            cert_p.space_after = Pt(2)
        doc.paragraphs[-1].space_after = Pt(6)

    # === EDUCATION ===
    if content.get('education'):
        add_section_heading(doc, 'Education')
        for edu in content['education']:
            edu_p = doc.add_paragraph()
            if isinstance(edu, dict):
                edu_p.add_run(edu['degree']).bold = True
                edu_p.add_run(f" - {edu.get('school', '')}")
            else:
                edu_p.add_run(edu)
            edu_p.space_after = Pt(4)

    doc.save(output_path)
    return output_path


# =============================================================================
# Main Entry Point
# =============================================================================

def generate_application_documents(candidate, job, cover_letter, resume, output_dir="/tmp/"):
    """
    Generate both cover letter and resume documents.

    Args:
        candidate: dict with candidate information
            - name: Full name
            - phone: Phone number
            - email: Email address
            - linkedin: LinkedIn URL
            - calendar: Calendar link (optional)
            - title: Target job title for resume header

        job: dict with job information
            - title: Job title
            - company: Company name
            - location: Job location
            - date: Letter date (optional, defaults to today)

        cover_letter: dict with cover letter content
            - opening: Opening paragraph text
            - sections: List of {title, paragraphs} dicts
            - closing: Closing text (optional)
            - signature_title: Title under signature

        resume: dict with resume content
            - summary: Summary paragraph
            - skills: List of {category, items} dicts
            - experience: List of job dicts with title, company, dates, bullets
            - certifications: List of cert strings or {name, detail} dicts
            - education: List of edu strings or {degree, school} dicts

        output_dir: Directory to save documents (default: /tmp/)

    Returns:
        Tuple of (cover_letter_path, resume_path)
    """
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    name_part = sanitize_filename(candidate['name'])
    role_part = sanitize_filename(job['title'])
    company_part = sanitize_filename(job['company'])

    cl_filename = f"{name_part}_Cover_Letter_{role_part}_{company_part}.docx"
    resume_filename = f"{name_part}_Resume_{role_part}_{company_part}.docx"

    cl_path = output_dir / cl_filename
    resume_path = output_dir / resume_filename

    create_cover_letter(candidate, job, cover_letter, str(cl_path))
    print(f"Cover letter saved: {cl_path}")

    create_resume(candidate, resume, str(resume_path))
    print(f"Resume saved: {resume_path}")

    return str(cl_path), str(resume_path)


# =============================================================================
# Web Chat Extras
# =============================================================================

def parse_uploaded_resume(file_path):
    """
    Extract text from an uploaded .docx resume.

    Uses Python's zipfile module to read the document XML — no external
    tools needed. Works in the Analysis sandbox.

    Args:
        file_path: Path to the uploaded .docx file

    Returns:
        str: Extracted text content, or None on error
    """
    try:
        with zipfile.ZipFile(file_path, 'r') as docx_zip:
            xml_content = docx_zip.read('word/document.xml').decode('utf-8')

        # Remove XML tags, preserving paragraph breaks
        text = re.sub(r'</w:p>', '\n', xml_content)
        text = re.sub(r'<[^>]+>', ' ', text)
        # Normalize whitespace within lines
        lines = [re.sub(r'[ \t]+', ' ', line).strip() for line in text.split('\n')]
        # Remove empty lines but preserve paragraph structure
        text = '\n'.join(line for line in lines if line)
        return text
    except zipfile.BadZipFile:
        print(f"Error: {file_path} is not a valid .docx file")
        return None
    except KeyError:
        print(f"Error: {file_path} does not contain expected Word document structure")
        return None
    except Exception as e:
        print(f"Error extracting text from {file_path}: {e}")
        return None


def generate_profile_yaml(candidate, qualifications, output_path="/tmp/profile.yaml"):
    """
    Generate a portable profile.yaml that can be reused across sessions.

    The profile contains candidate info and qualifications — everything
    needed to skip resume parsing on the next run.

    Args:
        candidate: dict with name, phone, email, linkedin, calendar
        qualifications: dict with summary, skills, experience,
                       certifications, education

    Returns:
        str: Path to the generated profile.yaml
    """
    # Import yaml or install it
    try:
        import yaml
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "pyyaml", "-q"])
        import yaml

    profile = {
        'candidate': {
            'name': candidate.get('name', ''),
            'phone': candidate.get('phone', ''),
            'email': candidate.get('email', ''),
            'linkedin': candidate.get('linkedin', ''),
            'calendar': candidate.get('calendar', ''),
        },
        'qualifications': qualifications,
    }

    with open(output_path, 'w', encoding='utf-8') as f:
        yaml.dump(profile, f, default_flow_style=False, allow_unicode=True, sort_keys=False)

    print(f"Profile saved: {output_path}")
    return output_path


def load_profile_yaml(file_path):
    """
    Load a previously saved profile.yaml.

    Args:
        file_path: Path to the profile.yaml file

    Returns:
        tuple: (candidate_dict, qualifications_dict) or (None, None) on error
    """
    try:
        import yaml
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "pyyaml", "-q"])
        import yaml

    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            profile = yaml.safe_load(f)

        candidate = profile.get('candidate', {})
        qualifications = profile.get('qualifications', {})
        return candidate, qualifications
    except Exception as e:
        print(f"Error loading profile: {e}")
        return None, None
