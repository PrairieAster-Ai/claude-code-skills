#!/usr/bin/env python3
"""
Generate styled Word documents for job applications.

This module creates ATS-friendly, professionally styled cover letters and resumes
in Microsoft Word format (.docx) with WCAG AA accessible colors.

Usage:
    from generate_word_docs import generate_application_documents

    generate_application_documents(
        candidate=candidate_data,
        job=job_data,
        cover_letter=cover_letter_content,
        resume=resume_content,
        output_dir="/path/to/output"
    )
"""

import sys
import platform

def check_dependencies():
    """Check that required dependencies are installed and provide helpful error messages."""
    missing = []

    try:
        import docx
    except ImportError:
        missing.append("python-docx")

    try:
        import yaml
    except ImportError:
        missing.append("pyyaml")

    if missing:
        print(f"ERROR: Missing required dependencies: {', '.join(missing)}")
        print()

        if platform.system() == "Darwin":  # macOS
            print("On macOS, use a virtual environment to install dependencies:")
            print()
            print("    python3 -m venv ~/.claude-venv")
            print("    source ~/.claude-venv/bin/activate")
            print(f"    pip install {' '.join(missing)}")
            print()
            print("Then add to your ~/.zshrc to auto-activate:")
            print("    source ~/.claude-venv/bin/activate")
        else:  # Linux and others
            print("Install with:")
            print(f"    pip3 install {' '.join(missing)} --user")

        print()
        sys.exit(1)

check_dependencies()

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
from pathlib import Path
import re
import yaml


# =============================================================================
# WCAG AA Accessible Colors (7:1+ contrast ratio on white)
# =============================================================================
ACCENT_COLOR = RGBColor(0x1a, 0x52, 0x76)  # Navy Blue #1a5276 - 8:1 contrast
TEXT_PRIMARY = RGBColor(0x1a, 0x20, 0x2c)  # Near-black #1a202c - 16:1 contrast
TEXT_SECONDARY = RGBColor(0x37, 0x41, 0x51)  # Dark Gray #374151 - 10:1 contrast

# Spacing constants
SPACE_BEFORE_HEADING = Pt(14)  # Space before section headings
SPACE_AFTER_HEADING = Pt(10)   # Padding under section headings (~15px)


# =============================================================================
# Helper Functions
# =============================================================================

def set_document_single_spacing(doc):
    """Set single line spacing as default for the entire document.

    Note: Only sets line spacing rule, not paragraph spacing (space_before/space_after).
    This allows individual paragraphs to control their own spacing while maintaining
    single line spacing within paragraphs.
    """
    style = doc.styles['Normal']
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    # Don't set space_after/space_before to 0 - let individual paragraphs control their spacing

    if 'List Bullet' in doc.styles:
        list_style = doc.styles['List Bullet']
        list_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def add_horizontal_line(paragraph):
    """Add a colored horizontal line below a paragraph."""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')  # 1.5pt line
    bottom.set(qn('w:color'), '1a5276')  # Accent color
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_section_heading(doc, text, is_first=False):
    """
    Add a styled section heading with proper spacing.

    Args:
        doc: Document object
        text: Heading text
        is_first: If True, reduces space before (for first section after header)

    Returns:
        The paragraph object
    """
    section = doc.add_paragraph()
    section.paragraph_format.space_before = Pt(6) if is_first else SPACE_BEFORE_HEADING
    section.paragraph_format.space_after = SPACE_AFTER_HEADING

    section_run = section.add_run(text)
    section_run.bold = True
    section_run.font.size = Pt(12)
    section_run.font.color.rgb = ACCENT_COLOR
    add_horizontal_line(section)

    return section


def sanitize_filename(text):
    """Convert text to a safe filename component."""
    # Remove or replace unsafe characters
    safe = re.sub(r'[^\w\s-]', '', text)
    safe = re.sub(r'[-\s]+', '_', safe)
    return safe.strip('_')


# =============================================================================
# Cover Letter Generator
# =============================================================================

def create_cover_letter(candidate, job, content, output_path):
    """
    Create a styled cover letter document.

    Args:
        candidate: dict with keys: name, phone, email, linkedin, calendar
        job: dict with keys: title, company, location, date (optional)
        content: dict with keys: opening, sections (list of {title, paragraphs}),
                 closing, signature_title
        output_path: Full path for output file
    """
    doc = Document()
    set_document_single_spacing(doc)

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # === HEADER ===
    # Name
    name_p = doc.add_paragraph()
    name_run = name_p.add_run(candidate['name'])
    name_run.bold = True
    name_run.font.size = Pt(22)
    name_run.font.color.rgb = ACCENT_COLOR
    name_p.space_after = Pt(2)

    # Contact info
    contact_parts = [candidate.get('phone', ''), candidate.get('email', ''),
                     candidate.get('linkedin', ''), candidate.get('calendar', '')]
    contact_text = ' | '.join(p for p in contact_parts if p)

    contact_p = doc.add_paragraph()
    contact_run = contact_p.add_run(contact_text)
    contact_run.font.size = Pt(10)
    contact_run.font.color.rgb = TEXT_SECONDARY
    add_horizontal_line(contact_p)
    contact_p.space_after = Pt(20)  # Space below contact info line

    # === DATE & RECIPIENT ===
    date_str = job.get('date', datetime.now().strftime('%B %d, %Y'))
    doc.add_paragraph(date_str).space_after = Pt(16)  # Space after date

    recipient = doc.add_paragraph()
    recipient.add_run(f"Hiring Manager\n{job['company']}\n{job.get('location', '')}".strip())
    recipient.space_after = Pt(12)

    # Subject line
    subject = doc.add_paragraph()
    subject_run = subject.add_run(f"Re: {job['title']}")
    subject_run.bold = True
    subject.space_after = Pt(16)  # Space after Re: line

    # Salutation
    doc.add_paragraph("Dear Hiring Manager,").space_after = Pt(10)

    # === OPENING PARAGRAPH ===
    opening = doc.add_paragraph()
    opening.add_run(content['opening'])
    opening.space_after = Pt(12)

    # === CONTENT SECTIONS ===
    for i, section in enumerate(content.get('sections', [])):
        add_section_heading(doc, section['title'], is_first=(i == 0))

        for para_content in section.get('paragraphs', []):
            if isinstance(para_content, dict):
                # Labeled paragraph: {label: "...", text: "...", highlights: [...]}
                p = doc.add_paragraph()
                if para_content.get('label'):
                    p.add_run(f"{para_content['label']}: ").bold = True

                text = para_content.get('text', '')
                highlights = para_content.get('highlights', [])
                _add_text_with_highlights(p, text, highlights)
                p.space_after = Pt(8)

            elif isinstance(para_content, list):
                # Bullet list
                for bullet_item in para_content:
                    bullet = doc.add_paragraph(style='List Bullet')
                    if isinstance(bullet_item, dict):
                        bullet.add_run(f"{bullet_item['label']}: ").bold = True
                        bullet.add_run(bullet_item.get('text', ''))
                    else:
                        bullet.add_run(bullet_item)
                    bullet.space_after = Pt(4)

            else:
                # Plain paragraph
                p = doc.add_paragraph()
                p.add_run(str(para_content))
                p.space_after = Pt(8)

    # === CLOSING ===
    closing_text = content.get('closing', 'Best regards,')
    doc.add_paragraph(closing_text).space_after = Pt(20)

    # Signature
    sig_name = doc.add_paragraph()
    sig_name.add_run(candidate['name']).bold = True
    sig_name.space_after = Pt(2)

    if content.get('signature_title'):
        sig_title = doc.add_paragraph()
        sig_run = sig_title.add_run(content['signature_title'])
        sig_run.font.color.rgb = TEXT_SECONDARY
        sig_run.font.size = Pt(10)

    # Save
    doc.save(output_path)
    return output_path


def _add_text_with_highlights(paragraph, text, highlights):
    """Add text to paragraph with certain phrases bolded."""
    if not highlights:
        paragraph.add_run(text)
        return

    # Sort highlights by length (longest first) to avoid partial matches
    highlights = sorted(highlights, key=len, reverse=True)

    # Simple approach: just add the text, bold highlights handled separately
    # For complex highlighting, would need regex splitting
    remaining = text
    for highlight in highlights:
        if highlight in remaining:
            parts = remaining.split(highlight, 1)
            if parts[0]:
                paragraph.add_run(parts[0])
            paragraph.add_run(highlight).bold = True
            remaining = parts[1] if len(parts) > 1 else ''

    if remaining:
        paragraph.add_run(remaining)


# =============================================================================
# Resume Generator
# =============================================================================

def create_resume(candidate, content, output_path):
    """
    Create a styled resume document.

    Args:
        candidate: dict with keys: name, title, phone, email, linkedin, calendar
        content: dict with keys: summary, skills (list of {category, items}),
                 experience (list of jobs), certifications, education
        output_path: Full path for output file
    """
    doc = Document()
    set_document_single_spacing(doc)

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # === HEADER ===
    # Name
    name_p = doc.add_paragraph()
    name_run = name_p.add_run(candidate['name'])
    name_run.bold = True
    name_run.font.size = Pt(24)
    name_run.font.color.rgb = ACCENT_COLOR
    name_p.space_after = Pt(2)

    # Title
    if candidate.get('title'):
        title_p = doc.add_paragraph()
        title_run = title_p.add_run(candidate['title'])
        title_run.bold = True
        title_run.font.size = Pt(12)
        title_p.space_after = Pt(4)

    # Contact info
    contact_parts = [candidate.get('phone', ''), candidate.get('email', ''),
                     candidate.get('linkedin', ''), candidate.get('calendar', '')]
    contact_text = ' | '.join(p for p in contact_parts if p)

    contact_p = doc.add_paragraph()
    contact_run = contact_p.add_run(contact_text)
    contact_run.font.size = Pt(10)
    contact_run.font.color.rgb = TEXT_SECONDARY
    contact_p.space_after = Pt(8)

    # === SUMMARY ===
    if content.get('summary'):
        add_section_heading(doc, 'Summary', is_first=True)
        summary_p = doc.add_paragraph()
        summary_p.add_run(content['summary'])
        summary_p.space_after = Pt(6)

    # === TECHNICAL SKILLS ===
    if content.get('skills'):
        add_section_heading(doc, 'Technical Skills')
        for skill in content['skills']:
            p = doc.add_paragraph()
            p.add_run(f"{skill['category']}:").bold = True
            p.add_run(f" {skill['items']}")
            p.space_after = Pt(4)
        doc.paragraphs[-1].space_after = Pt(6)

    # === PROFESSIONAL EXPERIENCE ===
    if content.get('experience'):
        add_section_heading(doc, 'Professional Experience')

        for idx, job in enumerate(content['experience']):
            # Job title - add space before to separate from previous job
            job_title_p = doc.add_paragraph()
            job_title_p.paragraph_format.space_before = Pt(0) if idx == 0 else Pt(12)
            job_title_run = job_title_p.add_run(job['title'])
            job_title_run.bold = True
            job_title_run.font.size = Pt(11)
            job_title_p.space_after = Pt(1)

            # Company and dates
            meta_p = doc.add_paragraph()
            meta_run = meta_p.add_run(f"{job['company']} | {job['dates']}")
            meta_run.font.color.rgb = TEXT_SECONDARY
            meta_run.font.size = Pt(10)
            meta_p.space_after = Pt(4)

            # Bullets - indented to group with job entry
            for bullet in job.get('bullets', []):
                bullet_p = doc.add_paragraph(style='List Bullet')
                bullet_p.paragraph_format.left_indent = Inches(0.35)

                if isinstance(bullet, dict):
                    # {text: "...", highlights: [...]}
                    _add_text_with_highlights(bullet_p, bullet['text'],
                                              bullet.get('highlights', []))
                else:
                    bullet_p.add_run(bullet)

                bullet_p.space_after = Pt(2)

    # === CERTIFICATIONS ===
    if content.get('certifications'):
        add_section_heading(doc, 'Certifications')
        for cert in content['certifications']:
            cert_p = doc.add_paragraph(style='List Bullet')
            if isinstance(cert, dict):
                cert_p.add_run(cert['name']).bold = True
                cert_p.add_run(f" - {cert.get('detail', '')}")
            else:
                cert_p.add_run(cert)
            cert_p.space_after = Pt(2)
        doc.paragraphs[-1].space_after = Pt(6)

    # === EDUCATION ===
    if content.get('education'):
        add_section_heading(doc, 'Education')
        for edu in content['education']:
            edu_p = doc.add_paragraph()
            if isinstance(edu, dict):
                edu_p.add_run(edu['degree']).bold = True
                edu_p.add_run(f" - {edu.get('school', '')}")
            else:
                edu_p.add_run(edu)
            edu_p.space_after = Pt(4)

    # Save
    doc.save(output_path)
    return output_path


# =============================================================================
# Main Entry Point
# =============================================================================

def generate_application_documents(candidate, job, cover_letter, resume, output_dir):
    """
    Generate both cover letter and resume documents.

    Args:
        candidate: dict with candidate information
            - name: Full name
            - phone: Phone number
            - email: Email address
            - linkedin: LinkedIn URL
            - calendar: Calendar link (optional)
            - title: Target job title for resume header

        job: dict with job information
            - title: Job title
            - company: Company name
            - location: Job location
            - date: Letter date (optional, defaults to today)

        cover_letter: dict with cover letter content
            - opening: Opening paragraph text
            - sections: List of {title, paragraphs} dicts
            - closing: Closing text (optional)
            - signature_title: Title under signature

        resume: dict with resume content
            - summary: Summary paragraph
            - skills: List of {category, items} dicts
            - experience: List of job dicts with title, company, dates, bullets
            - certifications: List of cert strings or {name, detail} dicts
            - education: List of edu strings or {degree, school} dicts

        output_dir: Directory to save documents

    Returns:
        Tuple of (cover_letter_path, resume_path)
    """
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    # Generate filenames
    name_part = sanitize_filename(candidate['name'])
    role_part = sanitize_filename(job['title'])
    company_part = sanitize_filename(job['company'])

    cl_filename = f"{name_part}_Cover_Letter_{role_part}_{company_part}.docx"
    resume_filename = f"{name_part}_Resume_{role_part}_{company_part}.docx"

    cl_path = output_dir / cl_filename
    resume_path = output_dir / resume_filename

    # Generate documents
    create_cover_letter(candidate, job, cover_letter, str(cl_path))
    print(f"Cover letter saved: {cl_path}")

    create_resume(candidate, resume, str(resume_path))
    print(f"Resume saved: {resume_path}")

    return str(cl_path), str(resume_path)


# =============================================================================
# Configuration Loading
# =============================================================================

def get_config_path():
    """Get the path to the config file."""
    return Path(__file__).parent / "config.yaml"


def get_example_config_path():
    """Get the path to the example config file."""
    return Path(__file__).parent / "config.example.yaml"


def load_config():
    """
    Load user configuration from config.yaml.

    Returns:
        dict: Configuration data, or None if config doesn't exist
    """
    config_path = get_config_path()
    if not config_path.exists():
        return None

    with open(config_path, 'r') as f:
        return yaml.safe_load(f)


def get_candidate_from_config(config=None):
    """
    Extract candidate information from config.

    Args:
        config: Config dict (will load from file if not provided)

    Returns:
        dict: Candidate information for document generation
    """
    if config is None:
        config = load_config()

    if config is None:
        return None

    candidate_config = config.get('candidate', {})
    return {
        'name': candidate_config.get('name', ''),
        'phone': candidate_config.get('phone', ''),
        'email': candidate_config.get('email', ''),
        'linkedin': candidate_config.get('linkedin', ''),
        'calendar': candidate_config.get('calendar', ''),
    }


def get_portfolio_projects(config=None):
    """
    Get portfolio projects from config.

    Args:
        config: Config dict (will load from file if not provided)

    Returns:
        list: Portfolio projects with achievements and technologies
    """
    if config is None:
        config = load_config()

    if config is None:
        return []

    return config.get('portfolio_projects', [])


def get_qualifications(config=None):
    """
    Get qualifications (resume data) from config.

    Args:
        config: Config dict (will load from file if not provided)

    Returns:
        dict: Qualifications including summary, skills, experience,
              certifications, and education. Returns None if not configured.
    """
    if config is None:
        config = load_config()

    if config is None:
        return None

    quals = config.get('qualifications', {})

    # Return None if qualifications section is empty or missing key data
    if not quals or not quals.get('experience'):
        return None

    return {
        'summary': quals.get('summary', ''),
        'skills': quals.get('skills', []),
        'experience': quals.get('experience', []),
        'certifications': quals.get('certifications', []),
        'education': quals.get('education', [])
    }


def get_output_dir(config=None):
    """
    Get output directory from config, with ~ expansion.

    Args:
        config: Config dict (will load from file if not provided)

    Returns:
        Path: Output directory path
    """
    if config is None:
        config = load_config()

    default_dir = "~/Documents/Job Application Docs/generated/"

    if config is None:
        return Path(default_dir).expanduser()

    output_dir = config.get('paths', {}).get('output_dir', default_dir)
    return Path(output_dir).expanduser()


# =============================================================================
# Example Usage / Test
# =============================================================================

if __name__ == "__main__":
    # Check if config exists
    config = load_config()

    if config:
        print("Configuration loaded from config.yaml")
        candidate = get_candidate_from_config(config)
        print(f"  Candidate: {candidate.get('name', 'Not set')}")
        print(f"  Output dir: {get_output_dir(config)}")

        projects = get_portfolio_projects(config)
        if projects:
            print(f"  Portfolio projects: {len(projects)}")
            for p in projects:
                print(f"    - {p.get('name')}")
    else:
        print("No config.yaml found.")
        print("Copy config.example.yaml to config.yaml and fill in your information.")
        print(f"  Example config: {get_example_config_path()}")

    print("\nImport this module to generate documents.")
    print("See generate_application_documents() for the main entry point.")
