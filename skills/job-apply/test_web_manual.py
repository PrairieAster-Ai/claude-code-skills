#!/usr/bin/env python3
"""
Manual test script for generate_word_docs_web.py

Run from the skill directory:
    python3 test_web_manual.py

Tests:
1. Document generation with sample data → cover letter + resume .docx
2. Profile YAML round-trip (with portfolio projects) → generate then load
3. Resume .docx parsing → extract text from generated resume
4. Highlight ordering → both highlights bolded regardless of position/length
5. Certification schema → both {detail} and {year, issuer} formats render correctly
6. Integer year in certification → str() coercion produces correct output
7. Empty-string highlights → no infinite loop, text preserved

All output goes to /tmp/job-apply-test/. Inspect the .docx files manually.
"""

import sys
import os
from pathlib import Path

# Import the web module from the same directory
sys.path.insert(0, str(Path(__file__).parent))
from generate_word_docs_web import (
    generate_application_documents,
    generate_profile_yaml,
    load_profile_yaml,
    parse_uploaded_resume,
    _add_text_with_highlights,
    _format_cert_detail,
)

OUTPUT_DIR = "/tmp/job-apply-test/"

# Sample data matching the schema from SKILL.md / generate_word_docs.py
SAMPLE_CANDIDATE = {
    "name": "Alex Rivera",
    "phone": "555.867.5309",
    "email": "alex.rivera@example.com",
    "linkedin": "linkedin.com/in/alexrivera",
    "calendar": "",
    "title": "Senior Software Engineer",
}

SAMPLE_JOB = {
    "title": "Staff Engineer",
    "company": "Acme Corp",
    "location": "Remote",
}

SAMPLE_COVER_LETTER = {
    "opening": (
        "With 8 years of full-stack experience and a track record of leading "
        "platform migrations serving millions of users, I'm excited to bring "
        "my expertise to the Staff Engineer role at Acme Corp."
    ),
    "sections": [
        {
            "title": "Technical Leadership",
            "paragraphs": [
                {
                    "label": "Platform Migration",
                    "text": (
                        "Led a team of 6 engineers through a 14-month migration "
                        "from monolith to microservices, reducing deployment time "
                        "from 4 hours to 12 minutes."
                    ),
                    "highlights": ["6 engineers", "4 hours to 12 minutes"],
                },
                [
                    "Designed event-driven architecture handling 50K req/sec",
                    "Mentored 3 junior engineers to mid-level promotions",
                ],
            ],
        },
        {
            "title": "Why Acme Corp",
            "paragraphs": [
                (
                    "Your commitment to developer experience aligns with my "
                    "passion for building tools that make engineering teams "
                    "more productive."
                ),
            ],
        },
    ],
    "closing": "I'd welcome the chance to discuss how I can contribute. Best regards,",
    "signature_title": "Senior Software Engineer",
}

SAMPLE_RESUME = {
    "summary": (
        "Full-stack engineer with 8 years of experience building scalable "
        "distributed systems. Led platform migration serving 2M+ users. "
        "Passionate about developer experience and engineering culture."
    ),
    "skills": [
        {"category": "Languages", "items": "Python, TypeScript, Go, Rust"},
        {"category": "Infrastructure", "items": "AWS, Kubernetes, Terraform, Docker"},
        {"category": "Databases", "items": "PostgreSQL, Redis, DynamoDB, Elasticsearch"},
    ],
    "experience": [
        {
            "title": "Senior Software Engineer",
            "company": "BigTech Inc.",
            "dates": "Mar 2021 - Present",
            "bullets": [
                {
                    "text": "Led team of 6 through monolith-to-microservices migration, reducing deploy time from 4 hours to 12 minutes",
                    "highlights": ["6", "4 hours to 12 minutes"],
                },
                {
                    "text": "Designed event-driven architecture processing 50K requests/second with 99.9% uptime",
                    "highlights": ["50K requests/second", "99.9% uptime"],
                },
                "Introduced trunk-based development, cutting release cycle from 2 weeks to daily",
            ],
        },
        {
            "title": "Software Engineer",
            "company": "StartupCo",
            "dates": "Jun 2018 - Feb 2021",
            "bullets": [
                {
                    "text": "Built real-time analytics pipeline processing 10M events/day",
                    "highlights": ["10M events/day"],
                },
                "Implemented CI/CD pipeline reducing build times by 60%",
            ],
        },
    ],
    # Use year/issuer format (config.yaml schema) to test _format_cert_detail
    "certifications": [
        {"name": "AWS Solutions Architect Professional", "year": "2023", "issuer": "Amazon"},
        {"name": "Certified Kubernetes Administrator", "year": "2022"},
    ],
    "education": [
        {"degree": "BS Computer Science", "school": "State University"},
    ],
}

SAMPLE_PORTFOLIO = [
    {
        "name": "AutoDeploy",
        "description": "CI/CD pipeline automation tool",
        "technologies": ["Python", "Docker", "GitHub Actions"],
        "achievements": [
            {"metric": "60% faster builds", "description": "Reduced CI pipeline time"},
        ],
    },
]

# Qualifications use the same cert format as SAMPLE_RESUME (year/issuer)
SAMPLE_QUALIFICATIONS = {
    "summary": SAMPLE_RESUME["summary"],
    "skills": SAMPLE_RESUME["skills"],
    "experience": SAMPLE_RESUME["experience"],
    "certifications": SAMPLE_RESUME["certifications"],
    "education": SAMPLE_RESUME["education"],
}


def test_document_generation():
    """Test 1: Generate cover letter and resume .docx files."""
    print("=" * 60)
    print("TEST 1: Document Generation")
    print("=" * 60)

    cl_path, resume_path = generate_application_documents(
        candidate=SAMPLE_CANDIDATE,
        job=SAMPLE_JOB,
        cover_letter=SAMPLE_COVER_LETTER,
        resume=SAMPLE_RESUME,
        output_dir=OUTPUT_DIR,
    )

    cl_ok = Path(cl_path).exists() and Path(cl_path).stat().st_size > 0
    resume_ok = Path(resume_path).exists() and Path(resume_path).stat().st_size > 0

    print(f"  Cover letter: {'PASS' if cl_ok else 'FAIL'} ({Path(cl_path).stat().st_size:,} bytes)")
    print(f"  Resume:       {'PASS' if resume_ok else 'FAIL'} ({Path(resume_path).stat().st_size:,} bytes)")
    print(f"  Inspect at:   {OUTPUT_DIR}")
    print()
    return cl_ok and resume_ok, resume_path


def test_profile_roundtrip():
    """Test 2: Generate profile.yaml with portfolio, then load it back."""
    print("=" * 60)
    print("TEST 2: Profile YAML Round-Trip (with portfolio)")
    print("=" * 60)

    profile_path = os.path.join(OUTPUT_DIR, "profile.yaml")
    generate_profile_yaml(
        SAMPLE_CANDIDATE, SAMPLE_QUALIFICATIONS,
        output_path=profile_path, portfolio_projects=SAMPLE_PORTFOLIO,
    )

    candidate, qualifications, portfolio = load_profile_yaml(profile_path)

    name_ok = candidate and candidate.get("name") == SAMPLE_CANDIDATE["name"]
    exp_ok = qualifications and len(qualifications.get("experience", [])) == 2
    skills_ok = qualifications and len(qualifications.get("skills", [])) == 3
    portfolio_ok = portfolio and len(portfolio) == 1 and portfolio[0]["name"] == "AutoDeploy"
    cert_ok = (qualifications and
               qualifications["certifications"][0].get("year") == "2023" and
               qualifications["certifications"][0].get("issuer") == "Amazon")

    print(f"  Name match:       {'PASS' if name_ok else 'FAIL'}")
    print(f"  Experience count:  {'PASS' if exp_ok else 'FAIL'}")
    print(f"  Skills count:      {'PASS' if skills_ok else 'FAIL'}")
    print(f"  Portfolio loaded:  {'PASS' if portfolio_ok else 'FAIL'}")
    print(f"  Cert year/issuer:  {'PASS' if cert_ok else 'FAIL'}")
    print(f"  Profile at:        {profile_path}")
    print()
    return name_ok and exp_ok and skills_ok and portfolio_ok and cert_ok


def test_resume_parsing(resume_path):
    """Test 3: Parse text from the generated resume .docx.

    Note: This is a circular test (parses our own output, not a real-world
    resume from Word/Google Docs). It verifies the extraction pipeline works
    but doesn't guarantee compatibility with all .docx producers.
    """
    print("=" * 60)
    print("TEST 3: Resume .docx Parsing (circular)")
    print("=" * 60)

    text = parse_uploaded_resume(resume_path)

    has_text = text is not None and len(text) > 100
    has_name = text is not None and "Alex Rivera" in text
    has_skill = text is not None and "Python" in text

    print(f"  Extracted text:  {'PASS' if has_text else 'FAIL'} ({len(text) if text else 0} chars)")
    print(f"  Contains name:   {'PASS' if has_name else 'FAIL'}")
    print(f"  Contains skills: {'PASS' if has_skill else 'FAIL'}")
    if text:
        preview = text[:200].replace("\n", " | ")
        print(f"  Preview: {preview}...")
    print()
    return has_text and has_name and has_skill


def test_highlight_ordering():
    """Test 4: _add_text_with_highlights handles out-of-order highlights.

    Regression test for the bug where highlights sorted by length would
    silently drop shorter highlights that appear earlier in the text.
    """
    print("=" * 60)
    print("TEST 4: Highlight Ordering")
    print("=" * 60)

    from docx import Document
    doc = Document()
    p = doc.add_paragraph()

    text = "We had 99.9% uptime processing 50K requests/second"
    highlights = ["50K requests/second", "99.9% uptime"]
    _add_text_with_highlights(p, text, highlights)

    # Collect all runs
    runs = [(r.text, r.bold) for r in p.runs]
    bold_texts = [r[0] for r in runs if r[1]]

    uptime_bolded = "99.9% uptime" in bold_texts
    reqs_bolded = "50K requests/second" in bold_texts
    full_text = "".join(r[0] for r in runs)
    text_preserved = full_text == text

    print(f"  '99.9% uptime' bolded:      {'PASS' if uptime_bolded else 'FAIL'}")
    print(f"  '50K requests/second' bolded: {'PASS' if reqs_bolded else 'FAIL'}")
    print(f"  Full text preserved:          {'PASS' if text_preserved else 'FAIL'}")
    print(f"  Runs: {runs}")
    print()
    return uptime_bolded and reqs_bolded and text_preserved


def test_cert_schema():
    """Test 5: _format_cert_detail handles both schema formats."""
    print("=" * 60)
    print("TEST 5: Certification Schema")
    print("=" * 60)

    # Old format: {name, detail}
    detail_cert = {"name": "AWS SAP", "detail": "Amazon, 2023"}
    detail_result = _format_cert_detail(detail_cert)
    detail_ok = detail_result == "Amazon, 2023"

    # New format: {name, year, issuer}
    year_cert = {"name": "AWS SAP", "year": "2023", "issuer": "Amazon"}
    year_result = _format_cert_detail(year_cert)
    year_ok = year_result == "Amazon, 2023"

    # Year only (no issuer)
    year_only_cert = {"name": "CKA", "year": "2022"}
    year_only_result = _format_cert_detail(year_only_cert)
    year_only_ok = year_only_result == "2022"

    # Empty cert (neither detail nor year)
    empty_cert = {"name": "Some Cert"}
    empty_result = _format_cert_detail(empty_cert)
    empty_ok = empty_result == ""

    print(f"  {{detail}} format:        {'PASS' if detail_ok else 'FAIL'} → '{detail_result}'")
    print(f"  {{year, issuer}} format:  {'PASS' if year_ok else 'FAIL'} → '{year_result}'")
    print(f"  {{year}} only format:     {'PASS' if year_only_ok else 'FAIL'} → '{year_only_result}'")
    print(f"  Empty cert:              {'PASS' if empty_ok else 'FAIL'} → '{empty_result}'")
    print()
    return detail_ok and year_ok and year_only_ok and empty_ok


def test_integer_year_cert():
    """Test 6: _format_cert_detail handles integer year without crashing.

    Regression test for the bug where integer years (from YAML without quotes)
    caused ', '.join() to fail with "expected str, got int".
    """
    print("=" * 60)
    print("TEST 6: Integer Year in Certification")
    print("=" * 60)

    # Integer year (YAML parsed without quotes)
    int_cert = {"name": "CKA", "year": 2023, "issuer": "CNCF"}
    int_result = _format_cert_detail(int_cert)
    int_ok = int_result == "CNCF, 2023"

    # Integer year only (no issuer)
    int_only = {"name": "CKA", "year": 2022}
    int_only_result = _format_cert_detail(int_only)
    int_only_ok = int_only_result == "2022"

    print(f"  Integer year + issuer:  {'PASS' if int_ok else 'FAIL'} → '{int_result}'")
    print(f"  Integer year only:      {'PASS' if int_only_ok else 'FAIL'} → '{int_only_result}'")
    print()
    return int_ok and int_only_ok


def test_empty_highlights():
    """Test 7: _add_text_with_highlights handles empty-string highlights safely.

    Regression test for the bug where an empty string in highlights
    would cause an infinite loop in str.find('', start).
    """
    print("=" * 60)
    print("TEST 7: Empty-String Highlights")
    print("=" * 60)

    from docx import Document
    doc = Document()
    p = doc.add_paragraph()

    text = "Reduced costs by 40% through optimization"
    highlights = ["40%", "", "  "]  # includes empty and whitespace-only
    _add_text_with_highlights(p, text, highlights)

    runs = [(r.text, r.bold) for r in p.runs]
    full_text = "".join(r[0] for r in runs)
    text_preserved = full_text == text
    bold_texts = [r[0] for r in runs if r[1]]
    pct_bolded = "40%" in bold_texts

    print(f"  Text preserved:   {'PASS' if text_preserved else 'FAIL'}")
    print(f"  '40%' bolded:     {'PASS' if pct_bolded else 'FAIL'}")
    print(f"  No crash:         PASS (reached this point)")
    print(f"  Runs: {runs}")
    print()
    return text_preserved and pct_bolded


def main():
    print()
    print("Job Apply Web Chat — Manual Test Suite")
    print("=" * 60)
    print(f"Output directory: {OUTPUT_DIR}")
    print()

    Path(OUTPUT_DIR).mkdir(parents=True, exist_ok=True)

    results = []

    ok, resume_path = test_document_generation()
    results.append(("Document Generation", ok))

    ok = test_profile_roundtrip()
    results.append(("Profile Round-Trip", ok))

    ok = test_resume_parsing(resume_path)
    results.append(("Resume Parsing (circular)", ok))

    ok = test_highlight_ordering()
    results.append(("Highlight Ordering", ok))

    ok = test_cert_schema()
    results.append(("Certification Schema", ok))

    ok = test_integer_year_cert()
    results.append(("Integer Year Cert", ok))

    ok = test_empty_highlights()
    results.append(("Empty-String Highlights", ok))

    # Summary
    print("=" * 60)
    print("SUMMARY")
    print("=" * 60)
    all_pass = True
    for name, passed in results:
        status = "PASS" if passed else "FAIL"
        print(f"  {status}  {name}")
        if not passed:
            all_pass = False

    print()
    if all_pass:
        print("All tests passed. Open the .docx files in Word to inspect formatting.")
    else:
        print("Some tests failed. Check output above.")

    print(f"\nOutput files in: {OUTPUT_DIR}")
    print()
    return 0 if all_pass else 1


if __name__ == "__main__":
    sys.exit(main())
